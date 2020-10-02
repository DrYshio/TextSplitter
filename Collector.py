import docx
import os


list_of_files = os.listdir(path=os.getcwd() + "\\texts")
sum_file = docx.Document()

for file in list_of_files:
    current_file = docx.Document(os.getcwd() + "\\texts\\" + str(file))
    for paragraph in current_file.paragraphs:
        sum_file.add_paragraph(paragraph.text)

sum_file.save('input.docx')