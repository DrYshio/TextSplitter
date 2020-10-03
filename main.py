import docx
import os


def merge_files(list_of_files):
    merged_file = docx.Document()
    for filename in list_of_files:
        current_file = docx.Document(os.path.join(os.getcwd(), 'texts', str(filename)))
        for paragraph in current_file.paragraphs:
            merged_file.add_paragraph(paragraph.text)

    merged_file.save('merged_output.docx')


def create_new_docx():
    document = docx.Document()
    return document


def number_of_symbols(input_docx):
    symbols_num = 0
    for paragraph in input_docx.paragraphs:
        symbols_num += len(paragraph.text)
    return symbols_num


try:
    os.mkdir(os.path.join(os.getcwd(), "texts"))
except FileExistsError:
    pass
try:
    os.mkdir(os.path.join(os.getcwd(), "output"))
except FileExistsError:
    pass

list_of_files = os.listdir(path=os.path.join(os.getcwd(), "texts"))
if len(list_of_files) > 1:
    merge_files(list_of_files)

try:
    sum_file = docx.Document('merged_output.docx')
except docx.opc.exceptions.PackageNotFoundError:
    sum_file = docx.Document('input.docx')


parts_num = input("Please specify the number of parts\n")
part_length = int(number_of_symbols(sum_file) / int(parts_num))
list_of_output = []
num = 0
i = 0
list_of_output.append(create_new_docx())


for paragraph in sum_file.paragraphs:
    list_of_output[i].add_paragraph(paragraph.text)
    num += len(paragraph.text)

    if num >= part_length:
        num = 0
        list_of_output[i].save(os.path.join(os.getcwd(), "output", f'{str(i)}.docx'))
        i += 1
        list_of_output.append(create_new_docx())

list_of_output[i].save(os.path.join(os.getcwd(), "output", f'{str(i)}.docx'))
